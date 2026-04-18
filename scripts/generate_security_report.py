"""Generate the Phase 1 security-audit DOCX report for the Agentic RAG System."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


SEVERITY_COLORS = {
    "HIGH": RGBColor(0xC0, 0x39, 0x2B),
    "MEDIUM": RGBColor(0xE6, 0x7E, 0x22),
    "LOW": RGBColor(0x27, 0x8B, 0x47),
    "INFO": RGBColor(0x2C, 0x3E, 0x50),
}


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 11, color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.text = k
        c2.text = v
        for run in c1.paragraphs[0].runs:
            run.bold = True
        c1.width = Cm(4.5)
        c2.width = Cm(12.0)


def add_finding(doc: Document, f: dict) -> None:
    heading = doc.add_paragraph()
    id_run = heading.add_run(f"{f['id']}  ")
    id_run.bold = True
    id_run.font.size = Pt(13)

    sev_run = heading.add_run(f"[{f['severity']}] ")
    sev_run.bold = True
    sev_run.font.size = Pt(13)
    sev_run.font.color.rgb = SEVERITY_COLORS[f["severity"]]

    title_run = heading.add_run(f["title"])
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Light List Accent 1"
    meta.autofit = True
    rows_data = [
        ("File / Line", f["location"]),
        ("OWASP / CWE", f["owasp"]),
        ("Severity", f["severity"]),
    ]
    for i, (k, v) in enumerate(rows_data):
        c1, c2 = meta.rows[i].cells
        c1.text = k
        c2.text = v
        for run in c1.paragraphs[0].runs:
            run.bold = True

    desc_p = doc.add_paragraph()
    desc_p.add_run("Description: ").bold = True
    desc_p.add_run(f["description"])

    fix_p = doc.add_paragraph()
    fix_run = fix_p.add_run("Fix: ")
    fix_run.bold = True
    fix_run.font.color.rgb = RGBColor(0x27, 0x8B, 0x47)
    fix_p.add_run(f["fix"])

    doc.add_paragraph()


def build_report(output_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Security Audit Report")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Agentic RAG System  —  Phase 1 (Static Analysis)")
    sub_run.italic = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    add_heading(doc, "Audit Metadata", level=2)
    add_kv_table(doc, [
        ("Repository", "Agentic_RAG_System"),
        ("Scope", "backend/api, backend/services, backend/utils, backend/core, frontend/src, rag-frontend/src"),
        ("Audit Type", "Static code review (Phase 1)"),
        ("Start Time", "02:05:53"),
        ("End Time", "02:10:12"),
        ("Total Elapsed", "4 min 19 s"),
        ("Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ])

    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "This report documents the findings of a Phase 1 static security audit of the "
        "Agentic RAG System. The system is a FastAPI + React Retrieval-Augmented Generation "
        "application that ingests PDF/CSV files into Pinecone and answers chat queries via "
        "OpenAI. The audit identified 18 findings: 5 High, 7 Medium, and 6 Low severity. "
        "The most critical issues cluster around missing authentication, unrestricted CORS, "
        "prompt-injection exposure, unbounded uploads, and a tenant-wipe path in the "
        "Pinecone deletion flow. No dangerous Python sinks (eval/exec/pickle/subprocess/"
        "shell=True) and no hardcoded API keys were found in tracked code.")

    add_heading(doc, "Severity Breakdown", level=2)
    sev_table = doc.add_table(rows=2, cols=4)
    sev_table.style = "Light Grid Accent 1"
    headers = ["High", "Medium", "Low", "Total"]
    values = ["5", "7", "6", "18"]
    for i, h in enumerate(headers):
        cell = sev_table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "1F3A5F")
    for i, v in enumerate(values):
        sev_table.rows[1].cells[i].text = v

    doc.add_paragraph()

    add_heading(doc, "Scope Notes", level=2)
    add_para(doc,
        "Several paths in the original checklist did not exist under the exact names requested. "
        "They were mapped to their real equivalents in the codebase:")
    mapping = [
        ("backend/api/chat.py", "backend/api/routes_chat.py"),
        ("backend/api/files.py", "backend/api/routes_files.py"),
        ("backend/services/file_service.py",
         "Not present - reviewed data_injestion_service.py + csv_processor.py"),
        ("backend/utils/phi_removal.py",
         "Not present - PHI logic lives in csv_processor.anonymize_data + medical_config.py"),
        ("frontend/src/components/Upload.tsx",
         "Not present - upload logic is inline in frontend/src/App.tsx and rag-frontend/src/App.tsx"),
    ]
    map_tbl = doc.add_table(rows=len(mapping) + 1, cols=2)
    map_tbl.style = "Light Grid Accent 1"
    hdr = map_tbl.rows[0].cells
    hdr[0].text = "Requested"
    hdr[1].text = "Actual in repo"
    for c in hdr:
        for run in c.paragraphs[0].runs:
            run.bold = True
    for i, (req, actual) in enumerate(mapping, start=1):
        map_tbl.rows[i].cells[0].text = req
        map_tbl.rows[i].cells[1].text = actual

    doc.add_paragraph()
    add_para(doc,
        "The Cursor 'AI Review' IDE action (Cmd+Shift+P) is a user-invoked command and "
        "cannot be triggered programmatically. The findings below come from equivalent "
        "static review: file-by-file reads, targeted regex scans, and a dependency check. "
        "Running 'AI Review' in the IDE manually is recommended as a complementary pass.",
        italic=True)

    doc.add_page_break()

    add_heading(doc, "2. Dangerous-Sink Scan Results", level=1)
    add_para(doc, "All of the following dangerous patterns were scanned across the "
             "backend and frontend (excluding venv and node_modules). No hits were found:")
    for item in [
        "eval(), exec(), __import__(), pickle.loads(), pickle.load()",
        "marshal.loads(), yaml.load()",
        "subprocess.*, os.system(), os.popen(), shell=True",
        "Direct SQL string concatenation",
        "dangerouslySetInnerHTML, innerHTML =, document.write()",
        "Hardcoded OpenAI (sk-...), Pinecone (pcsk_...) or Google (AIza...) keys in tracked files",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")

    add_para(doc, "Positive observations:", bold=True)
    for item in [
        "React interpolation auto-escapes output (no XSS sinks detected).",
        ".env is listed in .gitignore and not tracked in git.",
        "SSL trust store is explicitly pinned to certifi in core/config.py.",
        "Pydantic models validate all request bodies.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    add_heading(doc, "3. Findings", level=1)
    add_heading(doc, "3.1 High Severity", level=2)
    for f in HIGH_FINDINGS:
        add_finding(doc, f)

    add_heading(doc, "3.2 Medium Severity", level=2)
    for f in MEDIUM_FINDINGS:
        add_finding(doc, f)

    add_heading(doc, "3.3 Low Severity", level=2)
    for f in LOW_FINDINGS:
        add_finding(doc, f)

    doc.add_page_break()

    add_heading(doc, "4. Remediation Roadmap (Priority Order)", level=1)
    roadmap = [
        ("1", "Replace allow_origins=['*'] in backend/main.py:49 with an explicit origin allow-list."),
        ("2", "Add auth dependency (Depends(require_api_key)) to every router; enforce ownership inside delete/update paths."),
        ("3", "Constrain file_id: UUID4 and bound upload body size using MEDICAL_FILE_SETTINGS['max_file_size']."),
        ("4", "Swap PyPDF2 -> pypdf>=4.2 to eliminate known CVEs."),
        ("5", "Replace verbose HTTPException detail strings with opaque error IDs."),
        ("6", "Rebuild csv_processor.anonymize_data using presidio-analyzer; stop claiming HIPAA until audit log, AuthN, and encryption at rest are real."),
        ("7", "Add slowapi rate limiting and magic-byte file-type sniffing."),
        ("8", "Introduce prompt-injection guardrails (context delimiters, output moderation)."),
        ("9", "Move API_BASE into a Vite env var and require HTTPS in production builds."),
        ("10", "Stop persisting chat history (potential PHI) in localStorage."),
    ]
    roadmap_tbl = doc.add_table(rows=len(roadmap) + 1, cols=2)
    roadmap_tbl.style = "Light Grid Accent 1"
    hdr = roadmap_tbl.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Action"
    for c in hdr:
        for run in c.paragraphs[0].runs:
            run.bold = True
    for i, (n, action) in enumerate(roadmap, start=1):
        roadmap_tbl.rows[i].cells[0].text = n
        roadmap_tbl.rows[i].cells[1].text = action

    doc.add_paragraph()

    add_heading(doc, "5. Phase 2 Recommendations", level=1)
    for item in [
        "Run pip-audit and npm audit to cross-check dependency CVEs.",
        "Run Cursor 'AI Review' on backend/api, backend/services, backend/utils, frontend/src and reconcile with this report.",
        "Runtime probe of the unauthenticated endpoints using curl to confirm exploitability.",
        "Create real PHI fixtures and test csv_processor.anonymize_data for leakage.",
        "Add a CI step that runs bandit + semgrep rules on every pull request.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(output_path))


HIGH_FINDINGS = [
    {
        "id": "H-1",
        "severity": "HIGH",
        "title": "Wildcard CORS with credentials enabled",
        "location": "backend/main.py:47-53",
        "owasp": "A05:2021 Security Misconfiguration",
        "description": (
            "allow_origins=['*'] is combined with allow_credentials=True, allow_methods=['*'] "
            "and allow_headers=['*']. Any origin can call the API; if auth is later added via "
            "cookies the browser will also refuse the combination, signalling a broken posture."
        ),
        "fix": (
            "Replace ['*'] with an explicit allow-list (for example ['https://app.example.com']) "
            "and scope allow_methods / allow_headers to only what is needed."
        ),
    },
    {
        "id": "H-2",
        "severity": "HIGH",
        "title": "All endpoints unauthenticated / no authorization",
        "location": ("backend/api/routes_chat.py:35; "
                     "backend/api/routes_files.py:127, 221, 266, 344"),
        "owasp": "A01:2021 Broken Access Control; A07 Identification & Auth Failures",
        "description": (
            "POST /chat/, POST /files/add_file, DELETE /files/delete_file/{file_id}, "
            "PUT /files/update_file/{file_id} and POST /files/csv_info accept any caller. "
            "Any user can enumerate, delete or replace another tenant's vectors by file_id."
        ),
        "fix": (
            "Add a Depends(auth_dependency) that verifies a bearer token or API key and "
            "enforces per-file_id ownership before any mutation."
        ),
    },
    {
        "id": "H-3",
        "severity": "HIGH",
        "title": "Prompt injection (no input/output guardrails)",
        "location": ("backend/api/routes_chat.py:36-61; "
                     "backend/services/llm_service.py:99-115; "
                     "backend/services/rag_service.py:59-69"),
        "owasp": "OWASP LLM01 Prompt Injection; LLM02 Insecure Output Handling",
        "description": (
            "request.query is concatenated directly into the user-role prompt, and retrieved "
            "document text (itself attacker-uploadable via /files/add_file) is concatenated "
            "into the context with no delimiting, instruction-override defense, or output "
            "filter. A malicious uploaded PDF/CSV can override the system prompt at retrieval "
            "time (indirect prompt injection)."
        ),
        "fix": (
            "Introduce query sanitization, structured delimiters around retrieved context, "
            "an allow-list or moderation pass on outgoing answers, and a dedicated policy "
            "that instructs the model to ignore instructions found inside retrieved content."
        ),
    },
    {
        "id": "H-4",
        "severity": "HIGH",
        "title": "Unbounded file upload (DoS / memory exhaustion)",
        "location": "backend/api/routes_files.py:160-163, 302-305, 363-366",
        "owasp": "A04:2021 Insecure Design; A05 Misconfiguration",
        "description": (
            "await file.read() loads the entire request body into RAM before writing, with "
            "no Content-Length / max-size check. medical_config.MEDICAL_FILE_SETTINGS "
            "['max_file_size'] is defined but never referenced. Combined with H-2, an "
            "unauthenticated attacker can exhaust memory or disk."
        ),
        "fix": (
            "Stream the upload in chunks, enforce a max body size at the ASGI layer (for "
            "example via middleware), and compare against MEDICAL_FILE_SETTINGS['max_file_size']."
        ),
    },
    {
        "id": "H-5",
        "severity": "HIGH",
        "title": "Pinecone filter injection / unauthenticated tenant wipe",
        "location": ("backend/services/vectordb_service.py:143-170; "
                     "called from backend/api/routes_files.py:247"),
        "owasp": "A01 Broken Access Control; OWASP LLM08 Excessive Agency",
        "description": (
            "delete_by_file_id(file_id) builds filter={'file_id': {'$eq': file_id}} from an "
            "unvalidated path parameter (routes_files.py:222 only does .strip()). With no "
            "auth (H-2), any caller can guess or enumerate file_id values and delete "
            "vectors. Also top_k=10000 silently truncates larger result sets."
        ),
        "fix": (
            "Validate file_id against a strict UUID regex, require ownership proof before "
            "issuing the delete, and paginate the query instead of capping at 10k."
        ),
    },
]

MEDIUM_FINDINGS = [
    {
        "id": "M-1",
        "severity": "MEDIUM",
        "title": "Insecure temp-file handling",
        "location": "backend/api/routes_files.py:158-163, 300-305, 360-366",
        "owasp": "A04 Insecure Design; CWE-377",
        "description": (
            "tempfile.NamedTemporaryFile(delete=False, suffix=suffix) is correct for "
            "atomicity, but the suffix is derived from user-controlled file.filename, and "
            "if the await file.read() or write raises before the inner try/finally the "
            "temp file is orphaned. The file may also be world-readable on multi-user hosts."
        ),
        "fix": (
            "Wrap the whole block in try/finally, force a fixed suffix ('.pdf' / '.csv'), "
            "and call os.chmod(temp_file_path, 0o600) after creation."
        ),
    },
    {
        "id": "M-2",
        "severity": "MEDIUM",
        "title": "File-type validation is extension-only",
        "location": "backend/api/routes_files.py:145-156, 295, 356",
        "owasp": "A04 Insecure Design; CWE-434",
        "description": (
            "Only filename.lower().endswith(...) is checked. Content-Type and file magic "
            "bytes are not verified. A non-PDF payload renamed to .pdf is accepted up to "
            "the PyPDF2.PdfReader step, and the CSV path runs pandas.read_csv on arbitrary "
            "bytes which can trigger parser bugs."
        ),
        "fix": (
            "Sniff magic bytes (check for '%PDF-' or use python-magic) and validate "
            "UploadFile.content_type before saving."
        ),
    },
    {
        "id": "M-3",
        "severity": "MEDIUM",
        "title": "PyPDF2 is deprecated and has known DoS CVEs",
        "location": "backend/requirements.txt:11; backend/services/data_injestion_service.py:5, 42, 136",
        "owasp": "A06:2021 Vulnerable & Outdated Components",
        "description": (
            "PyPDF2 is end-of-life (upstream renamed to pypdf). CVE-2023-36464 (infinite "
            "loop) and other parsing DoS issues remain unpatched in PyPDF2 3.x."
        ),
        "fix": (
            "Swap to pypdf>=4.2: pip install pypdf, then change imports to "
            "'from pypdf import PdfReader'."
        ),
    },
    {
        "id": "M-4",
        "severity": "MEDIUM",
        "title": "HIPAA / PHI anonymization is naive and incomplete",
        "location": "backend/services/csv_processor.py:128-159",
        "owasp": "A02:2021 Cryptographic/Privacy Failures; HIPAA Safe Harbor 164.514(b)",
        "description": (
            "Only strips columns whose name matches a short regex list - free-text columns "
            "like clinical_notes or symptoms keep PHI. The bigram regex "
            "\\b[A-Z][a-z]+\\s+[A-Z][a-z]+\\b replaces medical terms like 'Diabetes Mellitus' "
            "or 'Heart Failure' with 'Patient', corrupting the corpus. No detection of SSN, "
            "DOB, phone, email, MRN, IP or ZIP within cell values. medical_config."
            "HIPAA_COMPLIANCE claims audit_logging, data_encryption and user_authentication "
            "= True but none of these are actually implemented."
        ),
        "fix": (
            "Use a real PHI detector (presidio-analyzer with HIPAA recognizers) on every "
            "string cell; separate de-identification from anonymization; implement real "
            "audit log, at-rest encryption, and AuthN before claiming HIPAA."
        ),
    },
    {
        "id": "M-5",
        "severity": "MEDIUM",
        "title": "Verbose internal-error disclosure to clients",
        "location": ("backend/api/routes_chat.py:80; "
                     "backend/api/routes_files.py:217, 262, 340, 408"),
        "owasp": "A09:2021 Security Logging & Monitoring Failures; CWE-209",
        "description": (
            "detail=f'Internal server error: {str(e)}' forwards raw backend exception "
            "strings (Pinecone stack frames, file paths, key names) to the unauthenticated "
            "HTTP client."
        ),
        "fix": (
            "Log full error server-side and return a generic opaque message with a "
            "correlation ID."
        ),
    },
    {
        "id": "M-6",
        "severity": "MEDIUM",
        "title": "CSV processing reads full file twice into memory",
        "location": "backend/services/csv_processor.py:170, 409",
        "owasp": "A04 Insecure Design (DoS)",
        "description": (
            "pd.read_csv(file_path, encoding=encoding) is called without nrows or chunksize, "
            "and again in get_csv_info. A large CSV will OOM the worker."
        ),
        "fix": (
            "Stream via chunksize=..., enforce an upper-bound row count before any full read."
        ),
    },
    {
        "id": "M-7",
        "severity": "MEDIUM",
        "title": "DEBUG env drives uvicorn --reload with no prod/dev separation",
        "location": "backend/core/config.py:43; backend/main.py:97-102",
        "owasp": "A05 Misconfiguration",
        "description": (
            "reload=settings.DEBUG means any deployer who exports DEBUG=true turns on "
            "file-watch reload - additional attack surface, race conditions, and stack "
            "traces in logs."
        ),
        "fix": (
            "Separate DEBUG from RELOAD; force reload=False unless APP_ENV == 'dev'."
        ),
    },
]

LOW_FINDINGS = [
    {
        "id": "L-1",
        "severity": "LOW",
        "title": "No rate limiting on any endpoint",
        "location": "backend/main.py (missing middleware)",
        "owasp": "A04 Insecure Design",
        "description": (
            "Combined with H-2, an attacker can freely drive OpenAI and Pinecone costs to "
            "infinity."
        ),
        "fix": "Add slowapi or an nginx/ingress rate-limit policy.",
    },
    {
        "id": "L-2",
        "severity": "LOW",
        "title": "Logging of queries / filenames may include PHI",
        "location": ("backend/api/routes_chat.py:50; "
                     "backend/api/routes_files.py:142, 197, 236; "
                     "backend/services/rag_service.py:107"),
        "owasp": "A09 Logging Failures; HIPAA audit log requirements",
        "description": (
            "logger.info(f'...{request.query[:100]}...') writes potentially PHI-laden "
            "queries to stdout."
        ),
        "fix": (
            "Redact or hash query text before logging; route audit logs to a HIPAA-"
            "compliant sink."
        ),
    },
    {
        "id": "L-3",
        "severity": "LOW",
        "title": "Hardcoded http:// backend URL in frontend (no env config)",
        "location": ("frontend/src/App.tsx:56; "
                     "rag-frontend/src/App.tsx:52, 111, 140-141"),
        "owasp": "A02 Cryptographic Failures (transport)",
        "description": (
            "const API_BASE = 'http://localhost:8000' is bundled as-is. Shipping this to "
            "production would send credentials and PHI over cleartext HTTP."
        ),
        "fix": (
            "Read from import.meta.env.VITE_API_BASE and require https:// in production builds."
        ),
    },
    {
        "id": "L-4",
        "severity": "LOW",
        "title": "Chat history (potentially PHI) persisted to localStorage",
        "location": "frontend/src/App.tsx:104-138",
        "owasp": "A02 Cryptographic Failures; A04 Insecure Design",
        "description": (
            "Any XSS or shared-kiosk reuse leaks every prior medical question and answer. "
            "localStorage is unencrypted and survives logout."
        ),
        "fix": (
            "Move to session-scoped, server-side chat history; clear on logout; avoid "
            "persisting assistant text verbatim."
        ),
    },
    {
        "id": "L-5",
        "severity": "LOW",
        "title": "file_id accepted as free-form string (no UUID enforcement)",
        "location": "backend/api/routes_files.py:222, 268",
        "owasp": "A03 Injection (defense-in-depth)",
        "description": (
            "Path parameter is only .strip()-checked. Not exploitable today because it is "
            "only used in a Pinecone filter, but feeds H-5."
        ),
        "fix": (
            "Declare file_id: UUID4 = Path(...); FastAPI will 422 non-UUIDs automatically."
        ),
    },
    {
        "id": "L-6",
        "severity": "LOW",
        "title": "No CSRF strategy documented",
        "location": "backend/main.py (no CSRF middleware)",
        "owasp": "A05 Misconfiguration",
        "description": (
            "Not exploitable today (no cookies or sessions), but the moment auth is added "
            "via cookies the API is vulnerable to CSRF."
        ),
        "fix": (
            "When adding auth, prefer Authorization: Bearer or add a double-submit CSRF token."
        ),
    },
]


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "docs" / "Security_Audit_Phase1.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    build_report(out)
    print(f"Wrote {out}")
